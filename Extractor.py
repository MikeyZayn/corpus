import re
import unicodedata
import pdfplumber
import docx
import argparse
import logging
import os
from datetime import datetime

# -------- Setup Logging --------
def setup_logger():
    log_dir = os.path.join("logs", "extractor")
    os.makedirs(log_dir, exist_ok=True)
    log_filename = os.path.join(log_dir, f"extractor_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log")

    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(message)s",
        handlers=[
            logging.FileHandler(log_filename, encoding="utf-8"),
            logging.StreamHandler()
        ]
    )
    logging.info(f"Logging initialized → {log_filename}")
    return log_filename

# -------- Extract text --------
def extract_text(file_path):
    if file_path.endswith(".pdf"):
        text = []
        with pdfplumber.open(file_path) as pdf:
            logging.info(f"Opened PDF with {len(pdf.pages)} pages")
            for i, page in enumerate(pdf.pages, start=1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
                        logging.info(f"Extracted page {i}/{len(pdf.pages)}")
                    else:
                        logging.warning(f"Page {i} had no extractable text")
                except Exception as e:
                    logging.error(f"Failed to extract page {i}: {e}")
        return "\n".join(text)

    elif file_path.endswith(".docx"):
        try:
            doc = docx.Document(file_path)
            logging.info(f"Opened DOCX with {len(doc.paragraphs)} paragraphs")
            return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        except Exception as e:
            logging.error(f"Failed to process DOCX: {e}")
            raise

    else:
        logging.error("Unsupported file type. Please use PDF or DOCX only.")
        raise ValueError("Unsupported file type. Please use PDF or DOCX only.")

# -------- Clean text --------
def clean_text(text):
    logging.info("Cleaning text (removing TOC, headers, page numbers)")
    text = re.sub(r'^\s*\d+\s*$', '', text, flags=re.MULTILINE)
    text = re.sub(r'(OKUQUKETHWE|TABLE OF CONTENTS).*', '', text, flags=re.IGNORECASE)
    text = re.sub(r'\n\s*\n', '\n', text)
    return text.strip()

# -------- Normalize --------
def normalize_text(text):
    logging.info("Normalizing text (Unicode NFKD)")
    return unicodedata.normalize("NFKD", text)

# -------- Sentence Split --------
def split_sentences(text):
    logging.info("Splitting text into sentences")
    sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', text) if s.strip()]
    logging.info(f"Generated {len(sentences)} sentences")
    return sentences

# -------- Save to TXT --------
def save_sentences(sentences, output_file):
    with open(output_file, "w", encoding="utf-8") as f:
        for s in sentences:
            f.write(s + "\n")
    logging.info(f"Saved sentences to {output_file}")

# -------- CLI Entry Point --------
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Extract sentences from PDF or DOCX into TXT.")
    parser.add_argument("input_file", help="Path to PDF or DOCX file")
    parser.add_argument("output_file", help="Path to output TXT file")
    args = parser.parse_args()

    log_file = setup_logger()

    try:
        raw_text = extract_text(args.input_file)
        cleaned = clean_text(raw_text)
        normalized = normalize_text(cleaned)
        sentences = split_sentences(normalized)

        save_sentences(sentences, args.output_file)
        logging.info(f"[✔] Extracted {len(sentences)} sentences → {args.output_file}")
    except Exception as e:
        logging.error(f"Extraction failed: {e}")

    logging.info(f"Run complete. Logs saved to {log_file}")
